#!/usr/bin/env python3
"""Read plain text from a Word .docx document using python-docx."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def read_document(path: str) -> str:
    docx_path = Path(path)
    if not docx_path.exists():
        raise FileNotFoundError(f"Word document not found: {docx_path}")

    document = Document(docx_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                lines.append(" | ".join(values))

    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Read text from a Word document.")
    parser.add_argument("path", help="Path to .docx file")
    args = parser.parse_args()
    print(read_document(args.path))


if __name__ == "__main__":
    main()
