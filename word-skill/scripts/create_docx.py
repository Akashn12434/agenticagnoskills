#!/usr/bin/env python3
"""Create a Word .docx document from structured content using python-docx."""

from __future__ import annotations

import argparse
import json
import os
import platform
import re
import subprocess
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = BASE_DIR / "output"


def slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", value.strip().lower()).strip("_")
    return slug or "document"

#Opens the generated document automatically.
def open_file(path: Path) -> None:
    """Open the generated document with the system default application."""
    if platform.system() == "Windows":
        os.startfile(path)  # type: ignore[attr-defined]
    elif platform.system() == "Darwin":
        subprocess.Popen(["open", str(path)])
    else:
        subprocess.Popen(["xdg-open", str(path)])


def unwrap_section(section: dict[str, Any]) -> dict[str, Any]:
    """Accept common nested model outputs and return the actual section dict."""
    if any(key in section for key in ("heading", "paragraphs", "bullets", "table")):
        return section

    for key in ("section", "content", "data", "example_key"):
        nested = section.get(key)
        if isinstance(nested, dict):
            return unwrap_section(nested)

    if len(section) == 1:
        nested = next(iter(section.values()))
        if isinstance(nested, dict):
            return unwrap_section(nested)

    return section


def default_sections(title: str) -> list[dict[str, Any]]:
    return [
        {
            "heading": "Introduction",
            "paragraphs": [
                f"This document provides a clear overview of {title}.",
                "It introduces the topic, explains the main ideas, and highlights practical relevance.",
            ],
        },
        {
            "heading": "Background",
            "paragraphs": [
                "Understanding the background helps place the topic in context.",
            ],
            "bullets": [
                "Key definitions and terms",
                "Important historical or current context",
                "Why the topic matters today",
            ],
        },
        {
            "heading": "Main Concepts",
            "paragraphs": [
                "The core concepts provide the foundation for understanding the topic.",
            ],
            "bullets": [
                "Primary idea with a concise explanation",
                "Supporting idea connected to the main topic",
                "Important implication for readers",
            ],
        },
        {
            "heading": "Conclusion",
            "paragraphs": [
                "The topic can be understood most clearly by connecting the background, main concepts, and practical impact.",
            ],
        },
    ]

#This is the cleanup stage.removes empty items, trims whitespace
def normalize_sections(title: str, sections: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    if not sections:
        return default_sections(title)

    normalized: list[dict[str, Any]] = []
    for index, section in enumerate(sections, start=1):
        section = unwrap_section(section)
        heading = str(section.get("heading") or section.get("title") or f"Section {index}").strip()
        paragraphs = [
            str(item).strip()
            for item in section.get("paragraphs", [])
            if str(item).strip()
        ]
        if not paragraphs and section.get("text"):
            paragraphs = [str(section["text"]).strip()]

        bullets = [
            str(item).strip()
            for item in section.get("bullets", [])
            if str(item).strip()
        ]
        table = section.get("table")
        normalized.append(
            {
                "heading": heading,
                "paragraphs": paragraphs,
                "bullets": bullets,
                "table": table if isinstance(table, dict) else None,
            }
        )

    return normalized

# This function applies consistent styles to the document, such as margins and fonts.
def apply_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11) #All regular text becomes Calibri 11pt for readability and a clean look.

    for style_name, size in (("Title", 22), ("Heading 1", 16), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)


def add_table(document: Document, table_data: dict[str, Any]) -> None:
    headers = [str(item) for item in table_data.get("headers", [])]
    rows = table_data.get("rows", [])
    if not headers or not rows:
        return

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for row in rows:
        values = row if isinstance(row, list) else [row.get(header, "") for header in headers]
        cells = table.add_row().cells
        for index, value in enumerate(values[: len(headers)]):
            cells[index].text = str(value)


def create_document(
    title: str,
    sections: list[dict[str, Any]] | None = None,
    subtitle: str | None = None,
    output_path: str | None = None,
    auto_open: bool = True,
) -> str:
    """Create a .docx file and return its absolute path."""
    normalized_sections = normalize_sections(title, sections)
    output = Path(output_path) if output_path else OUTPUT_DIR / f"{slugify(title)}.docx"
    if not output.is_absolute():
        output = OUTPUT_DIR / output

    output.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    apply_document_styles(document)

    title_paragraph = document.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        subtitle_paragraph = document.add_paragraph(subtitle)
        subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_paragraph = document.add_paragraph(date.today().isoformat())
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in normalized_sections:
        document.add_heading(section["heading"], level=1)
        for paragraph in section.get("paragraphs", []):
            document.add_paragraph(paragraph)
        for bullet in section.get("bullets", [])[:8]:
            document.add_paragraph(bullet, style="List Bullet")
        if section.get("table"):
            add_table(document, section["table"])

    document.save(output)
    resolved_output = output.resolve()
    if auto_open:
        open_file(resolved_output)
    return str(resolved_output)


def load_sections(value: str | None) -> list[dict[str, Any]] | None:
    if not value:
        return None
    path = Path(value)
    if path.exists():
        return json.loads(path.read_text(encoding="utf-8"))
    return json.loads(value)


def main() -> None:
    parser = argparse.ArgumentParser(description="Create a Word document.")
    parser.add_argument("title", help="Document title")
    parser.add_argument("--subtitle", help="Optional subtitle")
    parser.add_argument("--sections", help="JSON section list or path to JSON")
    parser.add_argument("--output", help="Output .docx path")
    parser.add_argument("--no-open", action="store_true", help="Do not open the document after creating it")
    args = parser.parse_args()

    result = create_document(
        args.title,
        load_sections(args.sections),
        args.subtitle,
        args.output,
        auto_open=not args.no_open,
    )
    print(result)


if __name__ == "__main__":
    main()
